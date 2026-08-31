"""Sinh 1 file .docx mau (gia lap de thi TN THPT) de kiem thu pipeline.

Chua: cong thuc OMML (dinh dang cong thuc Word native / MathType kieu
'Office compatible', duong chuyen doi chinh xac), 1 hinh anh (gia lap do
thi), va du 3 dang cau hoi: trac nghiem 4 dap an, dung/sai, tra loi ngan.
"""
import base64
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml

# PNG 1x1 trong suot, dung lam "hinh ve do thi" gia lap.
_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUBAScY42YAAAAASUVORK5CYII="
)

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _omml(inner: str) -> str:
    return f'<m:oMath xmlns:m="{M_NS}">{inner}</m:oMath>'


def _add_omath(paragraph, inner_xml: str):
    el = parse_xml(_omml(inner_xml))
    paragraph._p.append(el)


def build(path: str):
    doc = Document()

    doc.add_paragraph("SỞ GIÁO DỤC VÀ ĐÀO TẠO").alignment = 1
    doc.add_paragraph("ĐỀ THI THỬ TỐT NGHIỆP THPT - MÔN TOÁN").alignment = 1

    # ---- Cau 1: trac nghiem 4 dap an, co cong thuc phan so, dap an B in dam ----
    p = doc.add_paragraph("Câu 1: Cho biểu thức ")
    _add_omath(
        p,
        f'<m:f><m:fPr><m:ctrlPr/></m:fPr>'
        f'<m:num><m:r><m:t>x</m:t></m:r></m:num>'
        f'<m:den><m:r><m:t>2</m:t></m:r></m:den></m:f>'
        f'. Giá trị của biểu thức khi x=4 bằng bao nhiêu?',
    )
    doc.add_paragraph("A. 1")
    b = doc.add_paragraph()
    b.add_run("B. 2").bold = True
    doc.add_paragraph("C. 3")
    doc.add_paragraph("D. 4")

    img_path = Path("/tmp/_sample_graph.png")
    img_path.write_bytes(_PNG_1PX)

    # ---- Cau 2: TN co hinh ve (do thi), khong ve lai ma chen anh co san ----
    p2 = doc.add_paragraph("Câu 2: Cho đồ thị hàm số như hình vẽ sau:")
    doc.add_picture(str(img_path), width=None)
    doc.add_paragraph("Hàm số đồng biến trên khoảng nào sau đây?")
    a2 = doc.add_paragraph()
    a2.add_run("A. (-∞; 0)").bold = True
    doc.add_paragraph("B. (0; +∞)")
    doc.add_paragraph("C. (1; 2)")
    doc.add_paragraph("D. (-1; 1)")

    # ---- Cau 1 (Phan II): dung/sai 4 y ----
    p3 = doc.add_paragraph("Câu 1: Cho hàm số y = x")
    _add_omath(p3, '<m:sSup><m:e><m:r><m:t></m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>')
    p3.add_run(". Xét các khẳng định sau:")
    doc.add_paragraph("a) Hàm số xác định trên R. (Đúng)")
    doc.add_paragraph("b) Hàm số nghịch biến trên (0; +∞). (Sai)")
    doc.add_paragraph("c) Đồ thị hàm số đi qua gốc tọa độ. (Đúng)")
    doc.add_paragraph("d) Hàm số là hàm số lẻ. (Sai)")

    # ---- Cau 1 (Phan III): tra loi ngan ----
    p4 = doc.add_paragraph("Câu 1: Cho phương trình x")
    _add_omath(p4, '<m:sSup><m:e><m:r><m:t></m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>')
    p4.add_run(" - 4 = 0. Tính tổng các nghiệm của phương trình.")
    doc.add_paragraph("Đáp án: 0")

    doc.save(path)


if __name__ == "__main__":
    import sys
    build(sys.argv[1] if len(sys.argv) > 1 else "/tmp/sample_exam.docx")
    print("done")
